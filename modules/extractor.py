# ============================================
# MODULE 2: EXTRACTOR.PY
# Text Extraction — DOCX, TXT, PDF, Excel
# ============================================

import io
import os
import re
import sys
import tempfile
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import Colors, cprint
from modules.multilingual import clean_ocr_text, log_indicator

def extract_from_docx(file_path):
    try:
        import docx
        doc  = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        cprint(f"  ✅ DOCX: {len(text)} chars extracted", Colors.GREEN)
        return text.strip()
    except Exception as e:
        cprint(f"  ❌ DOCX Error: {e}", Colors.RED)
        return ""


def _is_weak_text_layer(text):
    cleaned = clean_ocr_text(text)
    signal_chars = len(re.findall(r"[A-Za-z0-9\u0900-\u097F\u0A80-\u0AFF]", cleaned))
    return signal_chars < 80 or len(cleaned.split()) < 12


def _read_uploaded_bytes(uploaded_file):
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)
    if hasattr(uploaded_file, "getvalue"):
        data = uploaded_file.getvalue()
    else:
        data = uploaded_file.read()
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)
    return data


def _extract_pdf_text_layer(pdf_bytes):
    import pdfplumber

    text = ""
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            extracted = page.extract_text() or ""
            if extracted.strip():
                text += f"\n[Page {index}]\n{extracted}\n"
    return clean_ocr_text(text)


def _ocr_pdf_pages(pdf_bytes):
    try:
        import pypdfium2 as pdfium
        from modules.ocr import perform_ocr_on_image
    except ImportError as exc:
        cprint(f"  PDF OCR warning: pypdfium2 is not installed ({exc}).", Colors.YELLOW)
        return ""

    tmp_path = None
    pages_text = []
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(pdf_bytes)
            tmp_path = tmp.name

        pdf = pdfium.PdfDocument(tmp_path)
        try:
            page_count = len(pdf)
            for page_index in range(page_count):
                page = pdf[page_index]
                try:
                    bitmap = page.render(scale=2.4)
                    image = bitmap.to_pil()
                    page_text = perform_ocr_on_image(image)
                    if page_text and page_text.strip():
                        pages_text.append(f"\n[OCR Page {page_index + 1}/{page_count}]\n{page_text.strip()}\n")
                finally:
                    try:
                        page.close()
                    except Exception:
                        pass
        finally:
            try:
                pdf.close()
            except Exception:
                pass
    except Exception as exc:
        cprint(f"  PDF OCR warning: could not render PDF pages ({exc}).", Colors.YELLOW)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    return clean_ocr_text("\n".join(pages_text))


def extract_text_from_file(uploaded_file):
    """
    Streamlit uploaded file se raw text nikalo.
    """
    text = ""
    ext = uploaded_file.name.lower().split(".")[-1]

    try:
        if ext == "docx":
            import docx

            doc = docx.Document(uploaded_file)

            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    row_data = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_data:
                        text += " | ".join(row_data) + "\n"

        elif ext == "pdf":
            pdf_bytes = _read_uploaded_bytes(uploaded_file)
            text_layer = _extract_pdf_text_layer(pdf_bytes)
            log_indicator("pdf_text_layer", text_length=len(text_layer))

            if text_layer and not _is_weak_text_layer(text_layer):
                text = text_layer
            else:
                cprint("  PDF text layer is empty or weak; trying OCR on rendered pages.", Colors.YELLOW)
                ocr_text = _ocr_pdf_pages(pdf_bytes)
                log_indicator("pdf_ocr", text_length=len(ocr_text))
                if ocr_text and len(ocr_text) > len(text_layer):
                    text = ocr_text
                elif text_layer and ocr_text:
                    text = clean_ocr_text(f"{text_layer}\n{ocr_text}")
                else:
                    text = text_layer or ocr_text

        elif ext in ["xlsx", "xls"]:
            import openpyxl

            wb = openpyxl.load_workbook(uploaded_file)
            sheet = wb.active

            headers = []
            first_row = True

            for row in sheet.iter_rows(values_only=True):
                if not any(row):
                    continue

                if first_row:
                    headers = [str(h) if h else "" for h in row]
                    first_row = False
                else:
                    for h, v in zip(headers, row):
                        if v is not None:
                            text += f"{h}: {v}\n"
                    text += "\n"

        elif ext == "csv":
            import pandas as pd

            df = pd.read_csv(uploaded_file)
            for _, row in df.iterrows():
                for column, value in row.items():
                    if pd.notna(value):
                        text += f"{column}: {value}\n"
                text += "\n"

        elif ext == "txt":
            text = uploaded_file.getvalue().decode("utf-8", errors="ignore")

        elif ext in ["jpg", "jpeg", "png", "bmp", "webp"]:
            from modules.ocr import extract_from_image

            tmp_path = None
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
                    tmp.write(uploaded_file.getbuffer())
                    tmp_path = tmp.name
                text = extract_from_image(tmp_path)
            finally:
                if tmp_path and os.path.exists(tmp_path):
                    try:
                        os.remove(tmp_path)
                    except OSError:
                        pass

        else:
            text = "Unsupported file format."

    except Exception as e:
        text = f"Error extracting text: {e}"

    return text

def extract_from_txt(file_path):
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read()
            cprint(f"  ✅ TXT ({encoding}): {len(text)} chars", Colors.GREEN)
            return text.strip()
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            cprint(f"  ❌ TXT Error: {e}", Colors.RED)
            return ""
    return ""

def extract_from_pdf(file_path):
    try:
        with open(file_path, "rb") as pdf_file:
            pdf_bytes = pdf_file.read()

        text_layer = _extract_pdf_text_layer(pdf_bytes)
        if text_layer and not _is_weak_text_layer(text_layer):
            cprint(f"  PDF: {len(text_layer)} chars extracted from text layer", Colors.GREEN)
            return text_layer.strip()

        cprint("  PDF has no usable text layer; trying OCR...", Colors.YELLOW)
        ocr_text = _ocr_pdf_pages(pdf_bytes)
        if ocr_text and len(ocr_text) > len(text_layer):
            return ocr_text.strip()
        if text_layer and ocr_text:
            return clean_ocr_text(f"{text_layer}\n{ocr_text}")
        return (text_layer or ocr_text).strip()

        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n[Page {i+1}/{total_pages}]\n" + page_text + "\n"
        if text.strip():
            cprint(f"  ✅ PDF: {len(text)} chars extracted", Colors.GREEN)
            return text.strip()
        cprint("  ⚠️  PDF has no text layer — trying OCR...", Colors.YELLOW)
        return ""
    except Exception as e:
        cprint(f"  ❌ PDF Error: {e}", Colors.RED)
        return ""

def extract_from_excel(file_path):
    try:
        import openpyxl
        wb   = openpyxl.load_workbook(file_path)
        text = ""
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text += f"\n[Sheet: {sheet_name}]\n"
            headers   = []
            first_row = True
            for row in sheet.iter_rows(values_only=True):
                if not any(row):
                    continue
                if first_row:
                    headers = [str(h) if h else "" for h in row]
                    text += " | ".join(headers) + "\n" + "-"*60 + "\n"
                    first_row = False
                else:
                    row_parts = [f"{h}: {v}" for h, v in zip(headers, row) if v is not None]
                    if row_parts:
                        text += "\n".join(row_parts) + "\n\n"
        cprint(f"  ✅ Excel: {len(text)} chars extracted", Colors.GREEN)
        return text.strip()
    except Exception as e:
        cprint(f"  ❌ Excel Error: {e}", Colors.RED)
        return ""
