# ============================================================
# CREATE_SAMPLES.PY — Saare Sample Documents Banao
# Alag-alag languages + formats mein!
# ============================================================

import os
from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors

BASE = os.path.dirname(os.path.abspath(__file__))
DOCS = os.path.join(BASE, "sample_documents")
os.makedirs(DOCS, exist_ok=True)

# ─────────────────────────────────────────────
# 1. ENGLISH Word Document
# ─────────────────────────────────────────────
def make_english_docx():
    doc = Document()
    doc.add_heading("Employee Information Form", 0)

    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    data = [
        ("First Name:",        "Pooja"),
        ("Surname:",           "Patel"),
        ("Father's Name:",     "Bhadhreshkumar Patel"),
        ("Employee ID:",       "DEV-402"),
        ("Department:",        "Software Development"),
        ("Role:",              "Frontend Engineer"),
        ("Date of Joining:",   "15-Aug-2021"),
        ("Email:",             "pooja.patel@company.com"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph()
    doc.add_paragraph(
        "This employee has demonstrated exceptional skill in frontend technologies "
        "including React and TypeScript. She leads the UI team for the SensorLive Dashboard "
        "project. Her performance in Q1 2025 was rated Outstanding by her manager."
    )
    path = os.path.join(DOCS, "english_pooja_patel.docx")
    doc.save(path)
    print(f"✅ Created: {os.path.basename(path)}")


# ─────────────────────────────────────────────
# 2. HINDI Word Document
# ─────────────────────────────────────────────
def make_hindi_docx():
    doc = Document()
    doc.add_heading("कर्मचारी सूचना प्रपत्र", 0)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    data = [
        ("पहला नाम (First Name):",   "Rahul"),
        ("उपनाम (Surname):",         "Sharma"),
        ("पिता का नाम (Father):",    "Rajesh Sharma"),
        ("कर्मचारी आईडी (ID):",      "EMP-301"),
        ("विभाग (Department):",      "UI/UX Design"),
        ("ईमेल (Email):",            "rahul.sharma@company.com"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph()
    doc.add_paragraph(
        "यह कर्मचारी UI/UX डिज़ाइन में विशेषज्ञ है। "
        "इन्होंने SensorLive Dashboard का पूरा डिज़ाइन तैयार किया है। "
        "This employee specializes in UI/UX design and has completed the full "
        "dashboard layout for the SensorLive project in Q1 2025."
    )
    path = os.path.join(DOCS, "hindi_rahul_sharma.docx")
    doc.save(path)
    print(f"✅ Created: {os.path.basename(path)}")


# ─────────────────────────────────────────────
# 3. GUJARATI Word Document
# ─────────────────────────────────────────────
def make_gujarati_docx():
    doc = Document()
    doc.add_heading("કર્મચારી માહિતી ફોર્મ", 0)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    data = [
        ("પ્રથમ નામ (First Name):",  "Hardik"),
        ("અટક (Surname):",           "Pandya"),
        ("પિતાનું નામ (Father):",    "Rameshbhai Pandya"),
        ("કર્મચારી ID (ID):",        "DEV-501"),
        ("વિભાગ (Department):",      "DevOps"),
        ("ઇમેઇલ (Email):",           "hardik.pandya@company.com"),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    doc.add_paragraph()
    doc.add_paragraph(
        "આ કર્મચારી DevOps ક્ષેત્રમાં નિષ્ણાત છે. "
        "This employee is an expert in DevOps and manages CI/CD pipelines "
        "for all projects at Pavitrasoft. He has automated 15+ deployment workflows in 2025."
    )
    path = os.path.join(DOCS, "gujarati_hardik_pandya.docx")
    doc.save(path)
    print(f"✅ Created: {os.path.basename(path)}")


# ─────────────────────────────────────────────
# 4. EXCEL Document
# ─────────────────────────────────────────────
def make_excel():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Employee Report"

    # Headers
    headers = ["First Name", "Surname", "Father's Name",
               "Employee ID", "Department", "Role",
               "Q1 Tasks", "Q1 Bugs Fixed", "Rating"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
        cell.fill = openpyxl.styles.PatternFill("solid", fgColor="1F4E79")

    # Data rows
    rows = [
        ("Amit",   "Verma",  "Anil Verma",    "DS-401", "Data Science",  "ML Engineer",  38, 5,  "Good"),
        ("Amit",   "Verma",  "Suresh Verma",  "FS-405", "Full Stack",    "Dev Lead",     52, 18, "Excellent"),
        ("Neha",   "Singh",  "Manoj Singh",   "BE-303", "Backend/IoT",   "IoT Architect",60, 8,  "Outstanding"),
        ("Kiran",  "Desai",  "Prakash Desai", "QA-201", "QA Testing",    "QA Lead",      75, 0,  "Outstanding"),
        ("Prince", "Patel",  "Rakeshbhai",    "MD-601", "Mobile Dev",    "Flutter Dev",  45, 10, "Good"),
    ]
    for r, row in enumerate(rows, 2):
        for c, val in enumerate(row, 1):
            ws.cell(row=r, column=c, value=val)

    path = os.path.join(DOCS, "employee_performance_report.xlsx")
    wb.save(path)
    print(f"✅ Created: {os.path.basename(path)}")


# ─────────────────────────────────────────────
# 5. PDF Document — English
# ─────────────────────────────────────────────
def make_pdf():
    path = os.path.join(DOCS, "neha_singh_report.pdf")
    c = canvas.Canvas(path, pagesize=A4)
    w, h = A4

    # Header bar
    c.setFillColor(colors.HexColor("#1F4E79"))
    c.rect(0, h - 80, w, 80, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 22)
    c.drawString(40, h - 50, "Employee Performance Report")
    c.setFont("Helvetica", 12)
    c.drawString(40, h - 68, "Pavitrasoft Pvt. Ltd. | Q1 2025")

    # Fields
    c.setFillColor(colors.black)
    fields = [
        ("First Name",    "Neha"),
        ("Surname",       "Singh"),
        ("Father's Name", "Manoj Singh"),
        ("Employee ID",   "BE-303"),
        ("Department",    "Backend / IoT"),
        ("Role",          "IoT Architect"),
        ("Email",         "neha.singh@company.com"),
        ("Q1 Tasks",      "60 Completed"),
        ("Rating",        "Outstanding"),
    ]
    y = h - 130
    for label, value in fields:
        c.setFont("Helvetica-Bold", 11)
        c.setFillColor(colors.HexColor("#1F4E79"))
        c.drawString(50, y, f"{label}:")
        c.setFont("Helvetica", 11)
        c.setFillColor(colors.black)
        c.drawString(220, y, value)
        y -= 28

    # Summary
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(colors.HexColor("#1F4E79"))
    c.drawString(50, y - 10, "Summary:")
    c.setFont("Helvetica", 10)
    c.setFillColor(colors.black)
    summary = [
        "Neha Singh is a highly skilled IoT Architect in the Backend team at Pavitrasoft.",
        "She successfully led the SensorLive Integration project, completing 60 tasks in Q1 2025",
        "with zero critical bugs. Her work has significantly improved system reliability by 40%.",
    ]
    y -= 35
    for line in summary:
        c.drawString(50, y, line)
        y -= 18

    c.save()
    print(f"✅ Created: {os.path.basename(path)}")


# ─────────────────────────────────────────────
# 6. TXT Document — French Language
# ─────────────────────────────────────────────
def make_french_txt():
    content = """Formulaire d'Information sur l'Employé
======================================
First Name: Priya
Surname: Shah
Father's Name: Dinesh Shah
Employee ID: HR-901
Department: Ressources Humaines (HR)
Role: Responsable RH Senior
Email: priya.shah@company.com

Résumé du Rapport:
Priya Shah est une professionnelle RH expérimentée avec plus de 5 ans d'expérience.
Elle gère le recrutement et le développement des talents pour toute l'organisation.
En Q1 2025, elle a recruté avec succès 12 nouveaux ingénieurs et lancé le programme
de formation interne. Sa performance est évaluée comme Excellente par la direction.

Note: This document is in French. The AI system will auto-detect and translate it.
"""
    path = os.path.join(DOCS, "french_priya_shah.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✅ Created: {os.path.basename(path)}")


# ─────────────────────────────────────────────
# 7. TXT Document — Spanish Language
# ─────────────────────────────────────────────
def make_spanish_txt():
    content = """Formulario de Información del Empleado
=======================================
First Name: Raj
Surname: Mehta
Father's Name: Sunil Mehta
Employee ID: FIN-101
Department: Finanzas (Finance)
Role: Analista Financiero Senior
Email: raj.mehta@company.com

Resumen del Informe:
Raj Mehta es un analista financiero altamente competente con experiencia en
contabilidad corporativa y análisis de inversiones. Durante el primer trimestre
de 2025, completó con éxito la auditoría anual y presentó el informe de presupuesto
al directorio. Su calificación de desempeño es Sobresaliente.

Note: This document is in Spanish. The AI system will auto-detect and translate it.
"""
    path = os.path.join(DOCS, "spanish_raj_mehta.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"✅ Created: {os.path.basename(path)}")


# ─────────────────────────────────────────────
# RUN ALL
# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("📁 Creating all sample documents...\n")
    make_english_docx()
    make_hindi_docx()
    make_gujarati_docx()
    make_excel()
    make_pdf()
    make_french_txt()
    make_spanish_txt()
    print(f"\n🎉 All sample documents created in: {DOCS}")
    print("\nFiles created:")
    for f in sorted(os.listdir(DOCS)):
        print(f"   📄 {f}")
