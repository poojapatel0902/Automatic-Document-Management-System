# ============================================================
# MAIN.PY — AI-Powered Document Management System
# Complete Pipeline:
# Ingest → Extract → Translate → NLP → Summarize → Route
# ============================================================

import os, sys, json, time
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from config import Colors, cprint, OUTPUT_DIR

# ============================================================
# CORE PIPELINE FUNCTION
# ============================================================

def process_single_document(file_path):
    """
    Ek document ke liye pura pipeline chalao!
    Step 1: Ingest  → file padho
    Step 2: Extract → text nikalo
    Step 3: Translate → English mein karo
    Step 4: NLP     → naam, ID nikalo
    Step 5: Summarize → summary banao
    Step 6: Route   → folder mein bhejo
    Returns: final JSON report
    """

    cprint("\n" + "█"*56, Colors.BOLD)
    cprint(f"  🤖 PROCESSING: {os.path.basename(file_path)}", Colors.BOLD)
    cprint("█"*56, Colors.BOLD)

    start_time = time.time()

    # ── STEP 1: INGEST ──────────────────────────────────
    cprint("\n[STEP 1/5] 📥 INGESTION", Colors.YELLOW)

    ext  = os.path.splitext(file_path)[1].lower()
    name = os.path.basename(file_path)

    if not os.path.exists(file_path):
        cprint(f"  ❌ File not found: {file_path}", Colors.RED)
        return None

    size_kb = os.path.getsize(file_path) / 1024
    file_info = {
        "file_name" : name,
        "file_path" : os.path.abspath(file_path),
        "extension" : ext,
        "file_type" : _get_type(ext),
        "size_label": f"{size_kb:.1f} KB",
    }
    cprint(f"  ✅ File: {name} ({size_kb:.1f} KB, type={file_info['file_type']})", Colors.GREEN)

    # ── STEP 2: EXTRACT TEXT ────────────────────────────
    cprint("\n[STEP 2/5] 📄 TEXT EXTRACTION", Colors.YELLOW)

    raw_text = _extract_text(file_path, ext)

    if not raw_text or not raw_text.strip():
        cprint("  ❌ No text extracted!", Colors.RED)
        return None
    cprint(f"  ✅ Extracted {len(raw_text)} characters", Colors.GREEN)

    # ── STEP 3: LANGUAGE DETECT + TRANSLATE ─────────────
    cprint("\n[STEP 3/5] 🌐 LANGUAGE & TRANSLATION", Colors.YELLOW)

    translated_text, original_lang, was_translated = _translate(raw_text)

    # ── STEP 4: NLP ENTITY EXTRACTION ───────────────────
    cprint("\n[STEP 4/5] 🧠 NLP ENTITY EXTRACTION", Colors.YELLOW)

    entities = _extract_entities(translated_text)
    _print_entities(entities)

    # ── STEP 5: SUMMARIZE ────────────────────────────────
    cprint("\n[STEP 5/5] 📝 SUMMARIZATION", Colors.YELLOW)

    summary = _summarize(translated_text)
    cprint(f"  Summary: {summary[:120]}...", Colors.CYAN)

    # ── ROUTE + SAVE ─────────────────────────────────────
    report = _route_and_save(
        file_info, entities, summary,
        original_lang, was_translated
    )

    elapsed = time.time() - start_time
    cprint(f"\n  ⏱️  Total time: {elapsed:.2f} seconds", Colors.CYAN)

    return report


# ============================================================
# HELPER — Get File Type
# ============================================================

def _get_type(ext):
    mapping = {
        ".docx": "document", ".txt": "document",
        ".pdf" : "pdf",
        ".xlsx": "spreadsheet", ".xls": "spreadsheet",
        ".jpg" : "image", ".jpeg": "image", ".png": "image",
    }
    return mapping.get(ext, "unknown")


# ============================================================
# HELPER — Extract Text by File Type
# ============================================================

def _extract_text(file_path, ext):
    try:
        if ext == ".docx":
            import docx
            doc  = docx.Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    parts = [c.text.strip() for c in row.cells if c.text.strip()]
                    if parts: text += "\n" + " | ".join(parts)
            return text

        elif ext == ".txt":
            for enc in ["utf-8","utf-16","latin-1"]:
                try:
                    with open(file_path,"r",encoding=enc) as f: return f.read()
                except UnicodeDecodeError: continue
            return ""

        elif ext == ".pdf":
            import pdfplumber
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t: text += t + "\n"
            return text

        elif ext in (".xlsx",".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(file_path)
            text = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                rows = list(ws.iter_rows(values_only=True))
                if not rows: continue
                headers = [str(h) for h in rows[0] if h]
                for row in rows[1:]:
                    if not any(row): continue
                    for h,v in zip(headers,row):
                        if v is not None: text += f"{h}: {v}\n"
                    text += "\n"
            return text

        elif ext in (".jpg",".jpeg",".png"):
            return _ocr_image(file_path)

    except Exception as e:
        cprint(f"  ❌ Extraction error: {e}", Colors.RED)
    return ""


def _ocr_image(file_path):
    """Image se OCR karo"""
    try:
        import pytesseract
        from PIL import Image, ImageEnhance, ImageFilter
        if os.name == "nt":
            pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        img = Image.open(file_path).convert("L")
        img = ImageEnhance.Contrast(img).enhance(2.0)
        img = img.filter(ImageFilter.SHARPEN)
        text = pytesseract.image_to_string(img, config="--psm 6")
        cprint(f"  ✅ OCR: {len(text)} chars", Colors.GREEN)
        return text
    except Exception as e:
        cprint(f"  ⚠️  OCR failed: {e}", Colors.YELLOW)
        return f"[Image file: {os.path.basename(file_path)} - OCR unavailable]"


# ============================================================
# HELPER — Translate
# ============================================================

def _translate(text):
    """Language detect + translate to English"""
    # Unicode-based language detection
    counts = {
        "hi": sum(1 for c in text if '\u0900' <= c <= '\u097F'),
        "gu": sum(1 for c in text if '\u0A80' <= c <= '\u0AFF'),
        "ar": sum(1 for c in text if '\u0600' <= c <= '\u06FF'),
        "zh": sum(1 for c in text if '\u4E00' <= c <= '\u9FFF'),
    }
    total    = max(len(text), 1)
    ratios   = {k: v/total for k,v in counts.items()}
    best     = max(ratios, key=ratios.get)
    detected = best if ratios[best] > 0.10 else "en"

    # French/German/Spanish detection via keywords
    if detected == "en":
        if any(w in text for w in ["Prénom","Département","société"]):   detected = "fr"
        elif any(w in text for w in ["Vorname","Abteilung","Mitarbeiter"]): detected = "de"
        elif any(w in text for w in ["Nombre","Apellido","Departamento"]): detected = "es"

    LANG_NAMES = {"en":"English","hi":"Hindi","gu":"Gujarati","mr":"Marathi",
                  "fr":"French","de":"German","es":"Spanish","ar":"Arabic","zh":"Chinese"}
    cprint(f"  🌐 Language detected: {LANG_NAMES.get(detected, detected)} ({detected})", Colors.CYAN)

    if detected == "en":
        cprint("  ✅ Already English — no translation needed", Colors.GREEN)
        return text, "en", False

    # Translate using deep-translator
    try:
        from deep_translator import GoogleTranslator
        chunks     = [text[i:i+4500] for i in range(0, len(text), 4500)]
        translator = GoogleTranslator(source=detected, target="en")
        result     = "\n".join(filter(None, [translator.translate(c) for c in chunks]))
        cprint(f"  ✅ Translated {detected} → English!", Colors.GREEN)
        return result, detected, True
    except Exception as e:
        cprint(f"  ⚠️  Translation skipped: {e}", Colors.YELLOW)
        return text, detected, False


# ============================================================
# HELPER — NLP Entity Extraction
# ============================================================

def _extract_entities(text):
    """Regex se saari entities nikalo"""
    import re

    def find(patterns, txt):
        for p in patterns:
            m = re.search(p, txt, re.IGNORECASE | re.MULTILINE)
            if m:
                return re.sub(r'\s+',' ', m.group(1)).strip()
        return None

    fn  = find([
        r"first\s*name\s*[:\-|]\s*([A-Za-z]+)",
        r"(?:employee|candidate|staff)\s*name\s*[:\-|]\s*([A-Za-z]+)",
        r"^name\s*[:\-|]\s*([A-Za-z]+)",
        r"vorname\s*[:\-|]\s*([A-Za-z]+)",
        r"pr[ée]nom\s*[:\-|]\s*([A-Za-z]+)",
        r"nombre\s*[:\-|]\s*([A-Za-z]+)",
    ], text)

    sn  = find([
        r"sur\s*name\s*[:\-|]\s*([A-Za-z]+)",
        r"last\s*name\s*[:\-|]\s*([A-Za-z]+)",
        r"family\s*name\s*[:\-|]\s*([A-Za-z]+)",
        r"nachname\s*[:\-|]\s*([A-Za-z]+)",
        r"nom\s*(?:de\s*famille)?\s*[:\-|]\s*([A-Za-z]+)",
        r"apellido\s*[:\-|]\s*([A-Za-z]+)",
    ], text)

    dad = find([
        r"father'?s?\s*name\s*[:\-|]\s*(.+?)(?:\n|$)",
        r"parent\s*name\s*[:\-|]\s*(.+?)(?:\n|$)",
        r"s/?o\s*[:\-|]?\s*(.+?)(?:\n|$)",
        r"vaters?\s*name\s*[:\-|]\s*(.+?)(?:\n|$)",
        r"nom\s*du\s*p[eè]re\s*[:\-|]\s*(.+?)(?:\n|$)",
    ], text)

    # Unique ID — Aadhar / PAN / Employee ID
    uid_type, uid_val = "Unknown", None
    aadhar = re.search(r'\b(\d{4}\s?\d{4}\s?\d{4})\b', text)
    pan    = re.search(r'\b([A-Z]{5}[0-9]{4}[A-Z])\b', text)
    empid  = re.search(r'(?:employee\s*id|emp\s*id|staff\s*id)\s*[:\-|]\s*([A-Z0-9\-]+)', text, re.I)

    if aadhar:  uid_type, uid_val = "Aadhar", aadhar.group(1)
    elif pan:   uid_type, uid_val = "PAN",    pan.group(1)
    elif empid: uid_type, uid_val = "Emp ID", empid.group(1)
    else:
        anyid = re.search(r'(?:id|no|number|#)\s*[:\-|]?\s*([A-Z0-9\-]{4,15})', text, re.I)
        if anyid: uid_type, uid_val = "ID", anyid.group(1)

    dob  = find([r"(?:dob|date\s*of\s*birth|born)\s*[:\-|]?\s*(\d{1,2}[-/]\w+[-/]\d{2,4})"], text)
    email= find([r"([a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z]{2,})"], text)
    phone= find([r"\b([6-9]\d{9})\b", r"\+\d[\d\s\-]{9,14}"], text)
    dept = find([r"(?:department|dept|division|abteilung|d[ée]partement)\s*[:\-|]\s*(.+?)(?:\n|$)"], text)

    return {
        "first_name" : fn,
        "surname"    : sn,
        "father_name": dad,
        "unique_id"  : {"type": uid_type, "value": uid_val},
        "dob"        : dob,
        "email"      : email,
        "phone"      : phone,
        "department" : dept,
    }


def _print_entities(e):
    uid = e.get("unique_id", {})
    print(f"""
  ┌─────────────────────────────────────────────┐
  │         📋 EXTRACTED ENTITIES               │
  ├─────────────────────────────────────────────┤
  │  First Name  : {str(e.get('first_name')  or '❌ Not found'):<28}│
  │  Surname     : {str(e.get('surname')     or '❌ Not found'):<28}│
  │  Father Name : {str(e.get('father_name') or '❌ Not found'):<28}│
  │  Unique ID   : {str((uid.get('type') or '?')+': '+(uid.get('value') or 'Not found')):<28}│
  │  DOB         : {str(e.get('dob')         or '❌ Not found'):<28}│
  │  Email       : {str(e.get('email')       or '❌ Not found'):<28}│
  │  Phone       : {str(e.get('phone')       or '❌ Not found'):<28}│
  │  Department  : {str(e.get('department')  or '❌ Not found'):<28}│
  └─────────────────────────────────────────────┘""")


# ============================================================
# HELPER — Summarize
# ============================================================

def _summarize(text):
    """Simple frequency-based summarization — always works!"""
    import re
    clean = re.sub(r'\[(?:Page|Sheet)[^\]]*\]|[-=|]{3,}', ' ', text)
    clean = re.sub(r'\s+', ' ', clean).strip()
    if len(clean) < 40:
        return clean

    sents = [s.strip() for s in re.split(r'(?<=[.!?])\s+', clean) if len(s.strip()) > 20]
    if len(sents) <= 3:
        return " ".join(sents)

    stop  = {"the","a","an","is","are","was","were","and","but","or","in",
             "on","at","to","of","for","by","with","this","that","it","he","she"}
    words = re.findall(r'\b[a-z]+\b', clean.lower())
    freq  = {}
    for w in words:
        if w not in stop and len(w) > 2:
            freq[w] = freq.get(w,0) + 1

    scores = {}
    for s in sents:
        score = sum(freq.get(w,0) for w in re.findall(r'\b[a-z]+\b', s.lower()))
        if re.search(r'(?i)(name|department|role|employee|performance|id)', s):
            score *= 1.5
        scores[s] = score

    top = sorted(scores, key=scores.get, reverse=True)[:3]
    return " ".join(s for s in sents if s in top)


# ============================================================
# HELPER — Route & Save
# ============================================================

def _route_and_save(file_info, entities, summary, lang, translated):
    """Folder banao, file copy karo, JSON save karo"""
    import shutil, re
    from datetime import datetime

    def safe(name):
        s = re.sub(r'[<>:"/\\|?*\n\r\t]','', str(name or "Unknown"))
        return re.sub(r'\s+','_', s.strip())[:40] or "Unknown"

    fn   = safe(entities.get("first_name")  or "Unknown")
    sn   = safe(entities.get("surname")     or "Unknown")
    dad  = safe(entities.get("father_name") or "Unknown")
    uid  = entities.get("unique_id",{})
    uid_v= safe(uid.get("value") or "NoID")

    folder = os.path.join(OUTPUT_DIR, fn, sn, f"{dad}_{uid_v}")
    os.makedirs(folder, exist_ok=True)

    # Copy file
    dest = os.path.join(folder, os.path.basename(file_info["file_path"]))
    if os.path.exists(dest):
        n,e  = os.path.splitext(os.path.basename(file_info["file_path"]))
        dest = os.path.join(folder, f"{n}_{datetime.now().strftime('%H%M%S')}{e}")
    shutil.copy2(file_info["file_path"], dest)

    # Build report
    report = {
        "timestamp"     : datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "file"          : file_info["file_name"],
        "file_type"     : file_info["file_type"],
        "original_lang" : lang,
        "was_translated": translated,
        "entities"      : {
            "first_name"  : entities.get("first_name"),
            "surname"     : entities.get("surname"),
            "father_name" : entities.get("father_name"),
            "unique_id"   : uid,
            "dob"         : entities.get("dob"),
            "email"       : entities.get("email"),
            "phone"       : entities.get("phone"),
            "department"  : entities.get("department"),
        },
        "summary"       : summary,
        "saved_to"      : dest,
        "folder"        : folder.replace(OUTPUT_DIR,"output").replace("\\","/"),
    }

    # Save JSON report
    with open(os.path.join(folder,"_report.json"),"w",encoding="utf-8") as f:
        json.dump(report, f, indent=2, ensure_ascii=False)

    # Print beautiful result
    _print_result(report)
    return report


def _print_result(r):
    e   = r["entities"]
    uid = e.get("unique_id",{})
    summ = r.get("summary","")
    folder_display = r.get("folder","")

    print(f"""
╔══════════════════════════════════════════════════════╗
║       ✅ DOCUMENT PROCESSED SUCCESSFULLY!            ║
╠══════════════════════════════════════════════════════╣
║  📄 File        : {str(r['file']):<35}║
║  🌐 Language    : {str(r['original_lang']):<35}║
║  🔄 Translated  : {str(r['was_translated']):<35}║
╠══════════════════════════════════════════════════════╣
║  👤 First Name  : {str(e.get('first_name')  or 'Not found'):<35}║
║  👤 Surname     : {str(e.get('surname')     or 'Not found'):<35}║
║  👤 Father Name : {str(e.get('father_name') or 'Not found'):<35}║
║  🪪 Unique ID   : {str((uid.get('type') or '?')+': '+(uid.get('value') or 'N/A')):<35}║
║  🏢 Department  : {str(e.get('department') or 'Not found'):<35}║
╠══════════════════════════════════════════════════════╣
║  📝 SUMMARY                                          ║""")
    for i in range(0, min(len(summ), 200), 52):
        print(f"║  {summ[i:i+52]:<52}  ║")
    print(f"""╠══════════════════════════════════════════════════════╣
║  📁 SAVED TO : {folder_display[:38]:<38}  ║
╚══════════════════════════════════════════════════════╝""")


# ============================================================
# PROCESS ENTIRE FOLDER
# ============================================================

def process_folder(folder_path):
    """Ek pore folder ke saare documents process karo!"""
    supported = (".docx",".txt",".pdf",".xlsx",".xls",".jpg",".jpeg",".png")

    files = [
        os.path.join(folder_path, f)
        for f in sorted(os.listdir(folder_path))
        if f.lower().endswith(supported)
    ]

    if not files:
        cprint(f"\n❌ No supported files found in: {folder_path}", Colors.RED)
        return []

    cprint(f"\n{'='*56}", Colors.BOLD)
    cprint(f"  📂 BATCH PROCESSING: {len(files)} documents", Colors.BOLD)
    cprint(f"  📁 Folder: {folder_path}", Colors.BOLD)
    cprint(f"{'='*56}", Colors.BOLD)

    reports = []
    success, failed = 0, 0

    for i, fpath in enumerate(files, 1):
        cprint(f"\n\n{'─'*56}", Colors.CYAN)
        cprint(f"  [{i}/{len(files)}] {os.path.basename(fpath)}", Colors.CYAN)
        cprint(f"{'─'*56}", Colors.CYAN)

        try:
            report = process_single_document(fpath)
            if report:
                reports.append(report)
                success += 1
            else:
                failed += 1
        except Exception as ex:
            cprint(f"\n  ❌ Error processing {os.path.basename(fpath)}: {ex}", Colors.RED)
            failed += 1

    # Batch summary
    cprint(f"\n\n{'█'*56}", Colors.BOLD)
    cprint(f"  🎉 BATCH COMPLETE!", Colors.BOLD)
    cprint(f"  ✅ Success : {success} files", Colors.GREEN)
    cprint(f"  ❌ Failed  : {failed} files", Colors.RED)
    cprint(f"  📁 Output  : {OUTPUT_DIR}", Colors.CYAN)
    cprint(f"{'█'*56}", Colors.BOLD)

    # Save batch summary JSON
    summary_path = os.path.join(OUTPUT_DIR, "_batch_summary.json")
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump({
            "total": len(files),
            "success": success,
            "failed": failed,
            "reports": reports,
        }, f, indent=2, ensure_ascii=False)
    cprint(f"\n  📊 Batch summary saved: {summary_path}", Colors.CYAN)

    return reports


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":
    BASE = os.path.dirname(os.path.abspath(__file__))
    SAMPLE_DOCS = os.path.join(BASE, "sample_documents")

    cprint("""
╔══════════════════════════════════════════════════════╗
║   🤖 AI-POWERED DOCUMENT MANAGEMENT SYSTEM          ║
║      Pavitrasoft Technologies — Demo v1.0           ║
╠══════════════════════════════════════════════════════╣
║  Features:                                           ║
║  ✅ Multi-format: DOCX, PDF, XLSX, TXT, Images      ║
║  ✅ Multi-language: EN, HI, GU, MR, FR, DE, ES      ║
║  ✅ Auto translation to English                      ║
║  ✅ Hierarchical Name Matching                       ║
║  ✅ Unique ID extraction (Aadhar/PAN/Employee)       ║
║  ✅ AI Summarization                                 ║
║  ✅ Auto folder routing + JSON reports               ║
╚══════════════════════════════════════════════════════╝
    """, Colors.CYAN)

    import argparse
    parser = argparse.ArgumentParser(description="AI Document Management System")
    parser.add_argument("--file",   help="Process a single file")
    parser.add_argument("--folder", help="Process all files in folder",
                        default=SAMPLE_DOCS)
    args = parser.parse_args()

    if args.file:
        process_single_document(args.file)
    else:
        process_folder(args.folder)
